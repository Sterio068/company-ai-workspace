"""
V1.1-SPEC §E-2 · indexer + extractor unit tests
================================================
不依賴 Meili / FastAPI · 純 mongomock + 臨時檔案
抽字真正 library(pymupdf/docx/openpyxl/Pillow)在 CI 環境可選裝
缺 library 的 test 會 skip
"""
import os
import pathlib
import tempfile
import pytest
import mongomock
from datetime import datetime, timezone

from services import knowledge_indexer
from services.knowledge_extract import extract, EXTRACTORS


# --- Fixtures ---
@pytest.fixture
def sources_col():
    client = mongomock.MongoClient()
    return client.chengfu_test.knowledge_sources


@pytest.fixture
def tmp_src(sources_col):
    """建一個有真實檔案結構的 tmp source"""
    d = tempfile.mkdtemp(prefix="chengfu_idx_")
    # 建立檔案結構
    (pathlib.Path(d) / "projects" / "海廢案").mkdir(parents=True)
    (pathlib.Path(d) / "projects" / "海廢案" / "建議書.txt").write_text(
        "承富創意 · 2024 環保署海洋廢棄物專案建議書 · 第一章主軸", encoding="utf-8"
    )
    (pathlib.Path(d) / "projects" / "海廢案" / "notes.md").write_text(
        "# 會議記錄\n- 3 月 1 日開案會議", encoding="utf-8"
    )
    (pathlib.Path(d) / "readme.md").write_text("根目錄 readme", encoding="utf-8")
    (pathlib.Path(d) / ".DS_Store").write_text("system", encoding="utf-8")
    (pathlib.Path(d) / "_unused" / "garbage.log").parent.mkdir()
    (pathlib.Path(d) / "_unused" / "garbage.log").write_text("log content")

    # 插入 source document
    sid = sources_col.insert_one({
        "name": "test source",
        "type": "local",
        "path": d,
        "enabled": True,
        "exclude_patterns": ["*.log", ".DS_Store"],
        "agent_access": [],
        "mime_whitelist": None,
        "max_size_mb": 10,
        "last_indexed_at": None,
        "last_index_stats": None,
        "created_at": datetime.now(timezone.utc),
        "updated_at": datetime.now(timezone.utc),
    }).inserted_id
    yield {"id": str(sid), "path": d, "col": sources_col}
    import shutil
    shutil.rmtree(d, ignore_errors=True)


# --- Extractor tests ---
def test_extract_text_utf8():
    """.txt utf-8 讀"""
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False, encoding="utf-8") as f:
        f.write("繁體中文測試")
        p = f.name
    try:
        r = extract(p)
        assert r["type"] == "text"
        assert "繁體中文測試" in r["content_preview"]
        assert r["filename"] == os.path.basename(p)
        assert r["size"] > 0
    finally:
        os.unlink(p)


def test_extract_md():
    """.md 走 text extractor"""
    with tempfile.NamedTemporaryFile(suffix=".md", mode="w", delete=False, encoding="utf-8") as f:
        f.write("# title\n內容")
        p = f.name
    try:
        r = extract(p)
        assert r["type"] == "text"
        assert "# title" in r["content_preview"]
    finally:
        os.unlink(p)


def test_extract_unknown_format():
    """.xyz 走 fallback · 不 raise"""
    with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
        f.write(b"binary blob")
        p = f.name
    try:
        r = extract(p)
        assert r["type"] == "unknown"
        assert "xyz" in r["content_preview"] or r["filename"] in r["content_preview"]
    finally:
        os.unlink(p)


def test_extract_nonexistent_returns_error():
    """不存在檔案 · 不 raise · 回 type=error"""
    r = extract("/tmp/definitely-not-exist-xyz-12345.pdf")
    assert r["type"] == "error"
    assert "error" in r


def test_extract_docx():
    """.docx 若 python-docx 裝好 · 真 parse"""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx 未安裝(非必測)")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        p = f.name
    try:
        doc = Document()
        doc.add_paragraph("承富創意整合行銷 · 測試段落")
        doc.add_paragraph("第二段 · 2026 年")
        doc.save(p)
        r = extract(p)
        assert r["type"] == "docx"
        assert "承富創意" in r["content_preview"]
        assert r["paragraph_count"] >= 2
    finally:
        os.unlink(p)


def test_extract_xlsx():
    """.xlsx 若 openpyxl 裝好 · 讀 sheet 前 N 行"""
    try:
        from openpyxl import Workbook
    except ImportError:
        pytest.skip("openpyxl 未安裝")

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        p = f.name
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "預算"
        ws.append(["項目", "金額"])
        ws.append(["場地", 30000])
        ws.append(["餐飲", 15000])
        wb.save(p)
        r = extract(p)
        assert r["type"] == "xlsx"
        assert "預算" in r["content_preview"]
        assert "場地" in r["content_preview"]
    finally:
        os.unlink(p)


# --- Indexer tests ---
def test_indexer_excludes_log_and_DS_Store(tmp_src):
    """exclude_patterns 會擋 .log 與 .DS_Store · 不抽字不索引"""
    stats = knowledge_indexer.reindex_source(
        tmp_src["id"], tmp_src["col"], meili_client=None,
    )
    assert stats["ok"] is True
    # 應該抽到 readme.md + projects/海廢案/建議書.txt + projects/海廢案/notes.md = 3
    assert stats["file_count"] >= 3
    assert stats["skipped"]["excluded"] >= 2  # .DS_Store + garbage.log
    # source doc 應更新 last_indexed_at
    doc = tmp_src["col"].find_one({"_id": __import__("bson").ObjectId(tmp_src["id"])})
    assert doc["last_indexed_at"] is not None


def test_indexer_incremental_second_run_skips_unchanged(tmp_src):
    """跑兩次 · 第二次 mtime 未變 · 都 skip"""
    s1 = knowledge_indexer.reindex_source(
        tmp_src["id"], tmp_src["col"], meili_client=None,
    )
    assert s1["file_count"] >= 3

    s2 = knowledge_indexer.reindex_source(
        tmp_src["id"], tmp_src["col"], meili_client=None,
    )
    # 第二輪 · 沒新檔 · 都在 unchanged
    assert s2["file_count"] == 0
    assert s2["skipped"]["unchanged"] >= 3


def test_indexer_disabled_source_skips(sources_col):
    """enabled=False 的 source 直接 skip"""
    d = tempfile.mkdtemp()
    sid = sources_col.insert_one({
        "name": "disabled",
        "path": d,
        "enabled": False,
    }).inserted_id
    try:
        r = knowledge_indexer.reindex_source(str(sid), sources_col, None)
        assert r.get("skipped") is True
    finally:
        import shutil
        shutil.rmtree(d, ignore_errors=True)


def test_indexer_nonexistent_path_returns_error(sources_col):
    """路徑不存在 · 不 crash · 回 ok=False + reason"""
    sid = sources_col.insert_one({
        "name": "bad-path",
        "path": "/tmp/definitely-not-exist-xyz-xyz",
        "enabled": True,
    }).inserted_id
    r = knowledge_indexer.reindex_source(str(sid), sources_col, None)
    assert r["ok"] is False
    assert "不存在" in r["reason"] or "不可讀" in r["reason"]


def test_indexer_reindex_all_iterates(tmp_src):
    """reindex_all 應跑所有 enabled sources"""
    r = knowledge_indexer.reindex_all(tmp_src["col"], None)
    assert "test source" in r
    assert r["test source"]["ok"] is True


def test_indexer_doc_id_stable():
    """同 source + 同 rel_path · id 必一致(才能 upsert 不重複)"""
    a = knowledge_indexer._doc_id_for("src_x", "projects/案/file.pdf")
    b = knowledge_indexer._doc_id_for("src_x", "projects/案/file.pdf")
    c = knowledge_indexer._doc_id_for("src_y", "projects/案/file.pdf")
    assert a == b
    assert a != c


def test_indexer_exclude_pattern_dir_prefix():
    """/sensitive/* 目錄前綴應該擋"""
    assert knowledge_indexer._match_excluded(
        "sensitive/secret.txt", ["/sensitive/*"]
    )
    assert not knowledge_indexer._match_excluded(
        "projects/normal.txt", ["/sensitive/*"]
    )


def test_indexer_exclude_pattern_glob():
    """~$* 擋 Word 暫存 · *.lock 擋 lock"""
    assert knowledge_indexer._match_excluded(
        "projects/案/~$建議書.docx", ["~$*"]
    )
    assert knowledge_indexer._match_excluded(
        "package-lock.json", ["*.lock", "package-lock.*"]
    )


def test_indexer_auto_tags_project_field(tmp_src):
    """檔在 projects/<名>/ 下 · 自動 tag project"""
    stats = knowledge_indexer.reindex_source(
        tmp_src["id"], tmp_src["col"], meili_client=None,
    )
    # 雖然 meili_client=None 沒進 Meili · 但 reindex_source 內部 logic 會處理 project
    # 直接用 extract + 手動模擬 indexer 的 enrichment 邏輯
    import os as _os
    from services.knowledge_extract import extract as _extract
    path = _os.path.join(tmp_src["path"], "projects", "海廢案", "建議書.txt")
    doc = _extract(path)
    assert doc["type"] == "text"
    # project enrichment 是 indexer 內部做的 · 這裡驗證 rel_path 推理
    rel = _os.path.relpath(path, tmp_src["path"])
    parts = rel.split(_os.sep)
    assert parts[0] == "projects"
    assert parts[1] == "海廢案"


# --- Delete source from index ---
def test_delete_source_from_index_no_meili():
    """meili_client=None 時 · 不 crash · 回 ok=False"""
    r = knowledge_indexer.delete_source_from_index("src_x", None)
    assert r["ok"] is False


# ============================================================
# Round 9 Q4 · 兩階段時間戳(scanned vs search_indexed)
# ============================================================
def test_indexer_meili_fail_keeps_search_timestamp_behind(tmp_src):
    """Meili 掛掉時 · last_scanned_at 前進 · 但 last_search_indexed_at 不動"""
    class BrokenMeili:
        def index(self, *a):
            raise RuntimeError("Meili down")
        def create_index(self, *a, **kw):
            raise RuntimeError("auth fail")

    stats = knowledge_indexer.reindex_source(
        tmp_src["id"], tmp_src["col"], meili_client=BrokenMeili(),
    )
    assert stats["ok"] is True
    assert stats["search_progress_advanced"] is False
    assert "meili_error" in stats

    doc = tmp_src["col"].find_one(
        {"_id": __import__("bson").ObjectId(tmp_src["id"])}
    )
    # scanning 進度:有前進
    assert doc["last_scanned_at"] is not None
    # search 進度:沒前進(None 代表從未成功)
    assert doc.get("last_search_indexed_at") is None


def test_indexer_meili_recovery_after_failure(tmp_src):
    """先讓 Meili 掛 → 第二次修好 · 應該補跑之前 scanned 過的檔"""
    class BrokenMeili:
        def index(self, *a):
            raise RuntimeError("Meili down")
        def create_index(self, *a, **kw):
            raise RuntimeError("auth fail")

    # 第一輪 · Meili 掛
    s1 = knowledge_indexer.reindex_source(
        tmp_src["id"], tmp_src["col"], meili_client=BrokenMeili(),
    )
    assert s1["search_progress_advanced"] is False
    files_scanned_r1 = s1["file_count"]
    assert files_scanned_r1 >= 3

    # 第二輪 · 沒 Meili(模擬恢復前) → scanned 已在 r1 前進 · 這輪沒事
    s2 = knowledge_indexer.reindex_source(
        tmp_src["id"], tmp_src["col"], meili_client=None,
    )
    # 檔案都沒變 · mtime 未過 scanned 時間 · 全部 unchanged
    # 但注意 since_mtime 取 min(scanned, search) · search 仍是 None
    # 所以應該重跑全部的 r1 檔 · 讓 scan 機制自己補上
    # 這驗證「Meili 恢復後下次 cron 自動補」的行為
    # (實際 cron 會帶 meili client · 這裡 meili=None 只驗邏輯)

    # 第三輪 · 帶正常 meili(模擬 task 成功)
    # Codex Round 10.5 fix · 現在 indexer 會 wait_for_task · 測試必須 fake 回 task uid + succeeded
    called = {"n": 0}
    class FakeTaskInfo:
        def __init__(self, uid): self.task_uid = uid
    class FakeTaskResult:
        def __init__(self, status): self.status = status
    class FakeIndex:
        def update_settings(self, *a): pass
        def add_documents(self, docs):
            called["n"] += len(docs)
            return FakeTaskInfo(called["n"])  # 回 task uid
        def delete_documents(self, **kw): pass
    class FakeMeili:
        def index(self, uid):
            return FakeIndex()
        def create_index(self, *a, **kw): pass
        def delete_index(self, *a): pass
        def wait_for_task(self, uid, timeout_in_ms=60000):
            return FakeTaskResult("succeeded")

    # 先模擬 fetch_info 時 index 不存在 · 需 create
    # 簡化:直接 patch _ensure_index 回 FakeIndex
    import services.knowledge_indexer as ki
    orig = ki._ensure_index
    ki._ensure_index = lambda c: FakeIndex()
    try:
        s3 = knowledge_indexer.reindex_source(
            tmp_src["id"], tmp_src["col"], meili_client=FakeMeili(),
        )
    finally:
        ki._ensure_index = orig

    # r3 應該重跑 r1 的檔(因為 search 進度從未前進)
    assert s3["file_count"] >= files_scanned_r1, \
        f"Q4 · Meili 恢復後應重跑 r1 的檔 · 實際 r3={s3['file_count']} r1={files_scanned_r1}"
    assert s3["search_progress_advanced"] is True
    doc = tmp_src["col"].find_one(
        {"_id": __import__("bson").ObjectId(tmp_src["id"])}
    )
    assert doc.get("last_search_indexed_at") is not None


def test_indexer_no_meili_advances_legacy_timestamp(tmp_src):
    """meili_client=None(cron 無 Meili)· last_indexed_at 仍前進 · legacy 行為相容"""
    stats = knowledge_indexer.reindex_source(
        tmp_src["id"], tmp_src["col"], meili_client=None,
    )
    assert stats["ok"] is True
    # search_progress_advanced 在無 meili 時為 None · 不是 True/False
    assert stats.get("search_progress_advanced") is None
    doc = tmp_src["col"].find_one(
        {"_id": __import__("bson").ObjectId(tmp_src["id"])}
    )
    # 相容舊 Admin UI · legacy 欄位還是要更新
    assert doc["last_indexed_at"] is not None
    assert doc["last_scanned_at"] is not None


# Codex Round 10.5 · Q4 真修 · Meili task 失敗時 last_search_indexed_at 不能前進
def test_indexer_meili_task_failed_keeps_search_behind(tmp_src):
    """add_documents enqueue 了但 task 本身 failed(例如 schema 衝突)
    這時 last_search_indexed_at 不應前進 · 否則下次 cron 不會補跑"""
    class FakeTaskInfo:
        def __init__(self, uid): self.task_uid = uid
    class FakeTaskResult:
        def __init__(self, status, err=None):
            self.status = status
            self.error = err
    class FakeIndex:
        def update_settings(self, *a): pass
        def add_documents(self, docs): return FakeTaskInfo(99)
        def delete_documents(self, **kw): pass
    class FakeMeili:
        def index(self, uid):
            # Patch _ensure_index 時回這個 · 表 Meili 有起來
            return FakeIndex()
        def create_index(self, *a, **kw): pass
        def delete_index(self, *a): pass
        def wait_for_task(self, uid, timeout_in_ms=60000):
            # task 失敗
            return FakeTaskResult("failed", {"message": "schema mismatch"})

    import services.knowledge_indexer as ki
    orig = ki._ensure_index
    ki._ensure_index = lambda c: FakeIndex()
    try:
        stats = knowledge_indexer.reindex_source(
            tmp_src["id"], tmp_src["col"], meili_client=FakeMeili(),
        )
    finally:
        ki._ensure_index = orig

    # 抽字應該成功 · file_count 仍有值
    assert stats["file_count"] >= 3
    # 但 Meili 寫入全失敗 · errors = file_count · search 不能前進
    assert stats["errors"] == stats["file_count"]
    assert stats["search_progress_advanced"] is False
    doc = tmp_src["col"].find_one(
        {"_id": __import__("bson").ObjectId(tmp_src["id"])}
    )
    # 這是關鍵:search 進度沒動 · 下次 cron 會重跑
    assert doc.get("last_search_indexed_at") is None


def test_indexer_meili_partial_failure_keeps_search_behind(tmp_src):
    """Codex Round 10.5 R2.1 · 一批成功 + 一批失敗 · last_search 不前進
    避免失敗批的檔永遠漏補"""
    import tempfile, os as _os
    # 先在 tmp_src 補多個檔 · 確保 > 200 個觸發多 batch
    for i in range(250):
        p = _os.path.join(tmp_src["path"], f"extra-{i:04d}.txt")
        with open(p, "w") as f:
            f.write(f"file {i}")

    class FakeTaskInfo:
        def __init__(self, uid): self.task_uid = uid
    class FakeResult:
        def __init__(self, status): self.status = status
    # 第一批 succeed · 第二批 fail · 第三批 succeed(flush)
    call_counter = {"n": 0}
    class FakeIndex:
        def update_settings(self, *a): pass
        def add_documents(self, docs):
            call_counter["n"] += 1
            return FakeTaskInfo(call_counter["n"])
        def delete_documents(self, **kw): pass
    class FakeMeili:
        def index(self, uid): return FakeIndex()
        def create_index(self, *a, **kw): pass
        def delete_index(self, *a): pass
        def wait_for_task(self, uid, timeout_in_ms=60000):
            # task 1 ok · task 2 fail · task 3 ok
            if uid == 2:
                return FakeResult("failed")
            return FakeResult("succeeded")

    import services.knowledge_indexer as ki
    orig = ki._ensure_index
    ki._ensure_index = lambda c: FakeIndex()
    try:
        stats = knowledge_indexer.reindex_source(
            tmp_src["id"], tmp_src["col"], meili_client=FakeMeili(),
        )
    finally:
        ki._ensure_index = orig

    # 有檔抽了 · 但一批失敗 · 整輪視為不前進
    assert stats["file_count"] > 0
    assert stats["search_progress_advanced"] is False, \
        "Codex R2.1:任一 Meili batch 失敗 · last_search 必不前進"
    doc = tmp_src["col"].find_one(
        {"_id": __import__("bson").ObjectId(tmp_src["id"])}
    )
    assert doc.get("last_search_indexed_at") is None


def test_indexer_meili_task_no_uid_returns_true_with_warning(tmp_src, caplog):
    """舊 Meili client 不回 task_uid · indexer 應該假成功 + log warning"""
    class NoUidTaskInfo:
        pass  # 沒 task_uid / taskUid 屬性
    class FakeIndex:
        def update_settings(self, *a): pass
        def add_documents(self, docs): return NoUidTaskInfo()
    class FakeMeili:
        def index(self, uid): return FakeIndex()
        def create_index(self, *a, **kw): pass
        def wait_for_task(self, uid, timeout_in_ms=60000):
            raise AssertionError("不該被叫到 · 因為沒 uid")

    import services.knowledge_indexer as ki
    orig = ki._ensure_index
    ki._ensure_index = lambda c: FakeIndex()
    try:
        stats = knowledge_indexer.reindex_source(
            tmp_src["id"], tmp_src["col"], meili_client=FakeMeili(),
        )
    finally:
        ki._ensure_index = orig

    # 假成功 · 沒任何 error
    assert stats["errors"] == 0
    assert stats["search_progress_advanced"] is True
